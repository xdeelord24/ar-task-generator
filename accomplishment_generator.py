"""
Accomplishment Report Generator
Generates professional accomplishment reports based on user tasks using HuggingFace API
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime, timedelta
import calendar
from typing import List, Dict
import os
import re

MONTH_NAME_PATTERN = (
    r'(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
)
MONTH_DATE_RANGE_PATTERN = re.compile(
    rf'(?P<month>{MONTH_NAME_PATTERN})\.?\s*'
    r'(?P<start>\d{1,2})'
    r'(?:\s*(?:-|–|to)\s*(?P<end>\d{1,2}))?'
    r'(?:\s*,?\s*(?P<year>\d{4}))?',
    re.IGNORECASE
)
NUMERIC_DATE_RANGE_PATTERN = re.compile(
    r'(?P<month>\d{1,2})[\/\-](?P<start>\d{1,2})'
    r'(?:\s*(?:-|–|to)\s*(?:(?P<end_month>\d{1,2})[\/\-])?(?P<end>\d{1,2}))?'
    r'(?:\s*,?\s*(?P<year>\d{4}))?',
    re.IGNORECASE
)
MONTH_MAP = {
    'jan': 1,
    'january': 1,
    'feb': 2,
    'february': 2,
    'mar': 3,
    'march': 3,
    'apr': 4,
    'april': 4,
    'may': 5,
    'jun': 6,
    'june': 6,
    'jul': 7,
    'july': 7,
    'aug': 8,
    'august': 8,
    'sep': 9,
    'sept': 9,
    'september': 9,
    'oct': 10,
    'october': 10,
    'nov': 11,
    'november': 11,
    'dec': 12,
    'december': 12,
}

try:
    from huggingface_hub import InferenceClient
    HF_HUB_AVAILABLE = True
except ImportError:
    HF_HUB_AVAILABLE = False
    InferenceClient = None

try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    ollama = None


class AccomplishmentGenerator:
    def __init__(self, huggingface_api_key: str = None, model_name: str = "meta-llama/Llama-3.1-8B-Instruct:novita", 
                 use_ollama: bool = False, ollama_model: str = "llama3.1", ollama_base_url: str = "http://localhost:11434"):
        """
        Initialize the generator
        
        Args:
            huggingface_api_key: HuggingFace API key for task enhancement (or set HF_TOKEN or HUGGINGFACE_API_KEY env var)
            model_name: HuggingFace model to use for text generation (default uses Llama-3.1-8B-Instruct:novita)
            use_ollama: If True, use Ollama for offline task enhancement
            ollama_model: Ollama model name (default: llama3.1)
            ollama_base_url: Ollama API base URL (default: http://localhost:11434)
        """
        self.api_key = huggingface_api_key or os.getenv('HF_TOKEN') or os.getenv('HUGGINGFACE_API_KEY')
        self.model_name = model_name
        self.client = None
        
        # Ollama configuration
        self.use_ollama = use_ollama
        self.ollama_model = ollama_model
        self.ollama_base_url = ollama_base_url
        
        if self.use_ollama and OLLAMA_AVAILABLE:
            try:
                # Initialize Ollama client
                if self.ollama_base_url != "http://localhost:11434":
                    self.ollama_client = ollama.Client(host=self.ollama_base_url)
                else:
                    self.ollama_client = ollama.Client()

                # Test Ollama connection - handle cloud models differently
                is_cloud_model = ':cloud' in self.ollama_model.lower() or 'cloud' in self.ollama_model.lower()
                if is_cloud_model:
                    # For cloud models, skip the list() test as it may not work
                    # Instead, try a simple generate call with minimal parameters
                    try:
                        test_response = self.ollama_client.generate(
                            model=self.ollama_model,
                            prompt="Hello",
                            options={'num_predict': 1}
                        )
                        print(f"OK Using Ollama cloud model: {self.ollama_model}")
                    except Exception as cloud_test_error:
                        print(f"Warning: Ollama cloud model '{self.ollama_model}' not available: {cloud_test_error}")
                        print("Falling back to HuggingFace or simple conversion")
                        self.use_ollama = False
                        self.ollama_client = None
                else:
                    # For local models, use the standard list() test
                    self.ollama_client.list()  # This will raise an error if Ollama is not running
                    print(f"OK Using Ollama with model: {self.ollama_model}")
            except Exception as e:
                print(f"Warning: Ollama not available: {e}")
                print("Falling back to HuggingFace or simple conversion")
                self.use_ollama = False
                self.ollama_client = None
        elif self.use_ollama and not OLLAMA_AVAILABLE:
            print("Warning: Ollama package not installed. Install with: pip install ollama")
            self.use_ollama = False
            self.ollama_client = None
        else:
            self.ollama_client = None
        
        if not self.use_ollama and self.api_key and HF_HUB_AVAILABLE:
            try:
                self.client = InferenceClient(api_key=self.api_key)
            except Exception as e:
                print(f"Warning: Could not initialize HuggingFace InferenceClient: {e}")
                self.client = None
        
    def calculate_date_ranges(self, year: int, month: int) -> tuple:
        """
        Calculate date ranges for the report period (1-15 or 15-end of month)
        Handles February and months with different day counts
        
        Args:
            year: Year of the report
            month: Month of the report (1-12)
            
        Returns:
            tuple: (start_date, end_date, is_first_half)
        """
        # Get the last day of the month
        last_day = calendar.monthrange(year, month)[1]
        
        # First half: 1-15
        start_first = datetime(year, month, 1)
        end_first = datetime(year, month, 15)
        
        # Second half: 16-last day
        start_second = datetime(year, month, 16)
        end_second = datetime(year, month, last_day)
        
        return (start_first, end_first, start_second, end_second, last_day)
    
    def get_working_days(self, start_date: datetime, end_date: datetime) -> List[datetime]:
        """
        Get list of working days (Monday-Friday) between two dates
        
        Args:
            start_date: Start date
            end_date: End date
            
        Returns:
            List of working day dates
        """
        working_days = []
        current_date = start_date
        
        while current_date <= end_date:
            # Monday = 0, Friday = 4
            if current_date.weekday() < 5:
                working_days.append(current_date)
            current_date += timedelta(days=1)
        
        return working_days
    
    def split_into_weeks(self, working_days: List[datetime]) -> List[Dict]:
        """
        Split working days into weeks (Monday-Friday)
        
        Args:
            working_days: List of working day dates
            
        Returns:
            List of week dictionaries with start and end dates
        """
        if not working_days:
            return []
        
        weeks = []
        current_week_start = working_days[0]
        current_week_end = working_days[0]
        
        for day in working_days[1:]:
            # If the day is more than 4 days after the week start, start a new week
            if (day - current_week_start).days > 4:
                weeks.append({
                    'start': current_week_start,
                    'end': current_week_end
                })
                current_week_start = day
                current_week_end = day
            else:
                current_week_end = day
        
        # Add the last week
        weeks.append({
            'start': current_week_start,
            'end': current_week_end
        })
        
        return weeks

    def _month_str_to_int(self, month_str: str, default_month: int) -> int:
        """
        Convert a month string to its numeric value, falling back to default_month.
        """
        if not month_str:
            return default_month
        key = month_str.strip().lower().rstrip('.')
        return MONTH_MAP.get(key, default_month)

    def _parse_task_date_ranges(self, task: str, default_year: int, default_month: int) -> List[Dict]:
        """
        Extract date ranges mentioned in a task string.
        Returns list of dictionaries with 'start' and 'end' datetime objects.
        """
        task = task or ""
        ranges = []

        # Month name formats (e.g., Nov 24-28, November 2)
        for match in MONTH_DATE_RANGE_PATTERN.finditer(task):
            month = self._month_str_to_int(match.group('month'), default_month)
            start_day = match.group('start')
            end_day = match.group('end') or start_day
            year = int(match.group('year')) if match.group('year') else default_year
            try:
                start_date = datetime(year, month, int(start_day))
                end_date = datetime(year, month, int(end_day))
                if start_date > end_date:
                    start_date, end_date = end_date, start_date
                ranges.append({'start': start_date, 'end': end_date})
            except ValueError:
                continue

        # Numeric formats (e.g., 11/24-11/28)
        for match in NUMERIC_DATE_RANGE_PATTERN.finditer(task):
            month = int(match.group('month'))
            start_day = match.group('start')
            end_month = match.group('end_month')
            end_day = match.group('end') or start_day
            year = int(match.group('year')) if match.group('year') else default_year
            if end_month:
                end_month_val = int(end_month)
            else:
                end_month_val = month
            try:
                start_date = datetime(year, month, int(start_day))
                end_date = datetime(year, end_month_val, int(end_day))
                if start_date > end_date:
                    start_date, end_date = end_date, start_date
                ranges.append({'start': start_date, 'end': end_date})
            except ValueError:
                continue

        return ranges

    def _find_week_index_for_range(self, date_range: Dict, weeks: List[Dict]) -> int:
        """
        Find the week index that overlaps with the provided date range.
        Returns -1 if no matching week is found.
        """
        if not date_range:
            return -1
        start = date_range['start'].date()
        end = date_range['end'].date()
        for idx, week in enumerate(weeks):
            week_start = week['start'].date()
            week_end = week['end'].date()
            if start <= week_end and end >= week_start:
                return idx
        return -1
    
    def enhance_task_with_ollama(self, task: str) -> str:
        """
        Enhance a task description using Ollama (offline mode)
        Converts to past tense and makes it more professional

        Args:
            task: Original task description

        Returns:
            Enhanced task in past tense
        """
        if not self.use_ollama or not OLLAMA_AVAILABLE or not self.ollama_client:
            return None

        try:
            # Use different prompts for cloud vs local models
            is_cloud_model = ':cloud' in self.ollama_model.lower() or 'cloud' in self.ollama_model.lower()

            if is_cloud_model:
                # Structured prompt for cloud models to get clean output
                prompt = f"""Task: {task}

Rewrite this task as one professional accomplishment statement in past tense. Start with a strong action verb. No first-person pronouns. Output only the accomplishment sentence, nothing else.

Example: "Developed a comprehensive project plan that streamlined workflow processes."

Accomplishment:"""
            else:
                # Original detailed prompt for local models
                prompt = (
                    "Rewrite the following task as a professional accomplishment in past tense. "
                    "Do not use first-person pronouns (I, we, my, our). "
                    "Start the sentence with a strong action verb and output only the rewritten accomplishment sentence(s) with no introductions or explanations. "
                    f"Task: {task}"
                )

            # Use Ollama generate API with adjusted parameters for cloud models
            generate_options = {
                'temperature': 0.3 if is_cloud_model else 0.5,  # Lower temperature for cloud models
                'num_predict': 200 if is_cloud_model else 150   # Higher limit for cloud models
            }

            response = self.ollama_client.generate(
                model=self.ollama_model,
                prompt=prompt,
                options=generate_options
            )

            # Handle different response formats for cloud vs local models
            if response:
                # Try different response field names
                enhanced_text = None
                for field in ['response', 'text', 'content', 'generated_text']:
                    if field in response and response[field]:
                        enhanced_text = response[field]
                        break

                # For cloud models, also check the thinking field
                if not enhanced_text and 'thinking' in response and response['thinking']:
                    thinking = response['thinking']
                    # Try to extract the final answer from thinking field
                    import re

                    # Look for quoted text that looks like accomplishment sentences
                    quotes = re.findall(r'"([^"]+)"', thinking)
                    if quotes:
                        # Filter for quotes that look like clean accomplishment sentences (not too long, start with action verb)
                        for quote in quotes:
                            quote = quote.strip()
                            if (len(quote.split()) >= 3 and len(quote.split()) <= 20 and
                                any(quote.lower().startswith(verb) for verb in
                                    ['developed', 'created', 'prepared', 'generated', 'authored',
                                     'delivered', 'produced', 'designed', 'implemented', 'compiled',
                                     'drafted', 'formulated', 'analyzed', 'conducted', 'performed',
                                     'completed', 'built', 'constructed', 'established', 'organized'])):
                                # Clean up any markdown or extra formatting
                                quote = re.sub(r'\*\*.*?\*\*', '', quote).strip()
                                # Make sure it doesn't contain colons or extra content
                                if ':' not in quote and '�' not in quote:
                                    enhanced_text = quote
                                    break

                    # If no good quotes found, try to find complete sentences
                    if not enhanced_text:
                        # Look for patterns like **Professional Accomplishment:** "text"
                        accomplishment_match = re.search(r'\*\*.*?\*\*\s*[""]([^""]+)[""]', thinking, re.IGNORECASE)
                        if accomplishment_match:
                            enhanced_text = accomplishment_match.group(1).strip()

                    # Last resort: find any sentence that looks like an accomplishment
                    if not enhanced_text:
                        sentences = re.split(r'[.!?]+', thinking)
                        for sentence in sentences:
                            sentence = sentence.strip()
                            # Skip sentences that are too short or contain instructions
                            if (len(sentence.split()) >= 4 and len(sentence.split()) <= 25 and
                                not any(word in sentence.lower() for word in ['thus', 'need', 'should', 'must', 'let me', 'example']) and
                                any(sentence.lower().startswith(verb) for verb in
                                    ['developed', 'created', 'prepared', 'generated', 'authored',
                                     'delivered', 'produced', 'designed', 'implemented'])):
                                enhanced_text = sentence
                                break

                if enhanced_text:
                    enhanced = self._clean_generated_text(enhanced_text)
                    if enhanced:
                        return enhanced

        except Exception as e:
            print(f"Error enhancing task with Ollama: {e}")
            # For cloud models, don't disable Ollama entirely on first error
            # as they might be more sensitive to network issues
            is_cloud_model = ':cloud' in self.ollama_model.lower() or 'cloud' in self.ollama_model.lower()
            if not is_cloud_model:
                # If Ollama fails for local models, disable it for this session
                self.use_ollama = False
                self.ollama_client = None

        return None
    
    def enhance_task_with_huggingface(self, task: str) -> str:
        """
        Enhance a task description using HuggingFace InferenceClient
        Converts to past tense and makes it more professional
        
        Args:
            task: Original task description
            
        Returns:
            Enhanced task in past tense
        """
        if not self.api_key or not self.client:
            # Fallback: Simple past tense conversion if no API key or client
            return None
        
        try:
            # Create a prompt for task enhancement
            prompt = (
                "Rewrite the following task as a professional accomplishment in past tense. "
                "Do not use first-person pronouns (I, we, my, our). "
                "Start the sentence with a strong action verb and output only the rewritten accomplishment sentence(s) with no introductions or explanations. "
                f"Task: {task}"
            )
            
            # Use InferenceClient chat completions
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=150,
                temperature=0.5,
            )
            
            if response and hasattr(response, 'choices') and len(response.choices) > 0:
                enhanced = self._clean_generated_text(response.choices[0].message.content)
                if enhanced:
                    return enhanced
                
        except Exception as e:
            print(f"Error enhancing task with HuggingFace: {e}")
        
        return None

    def _clean_generated_text(self, text: str) -> str:
        """
        Clean AI-generated text by removing introductory phrases or bullet markers.
        """
        if not text:
            return text
        cleaned = text.strip().strip('`"')
        # Remove leading bullet markers like "- " or "* "
        cleaned = re.sub(r'^[-*]\s+', '', cleaned)
        # Remove common introductory phrases
        intro_patterns = [
            r'^here\s+(?:is|are|was|were)\s+(?:an?|the)?\s*(?:rewritten|revised|updated)?\s*(?:version\s+of\s+)?(?:the\s+)?(?:task|accomplishment|sentence)\s*:?\s*',
            r'^(?:rewritten|revised)\s+(?:task|version)\s*:?\s*',
            r'^(?:professional\s+accomplishment|accomplishment)\s*:?\s*',
            r'^the\s+following\s+is\s*:?\s*',
        ]
        for pattern in intro_patterns:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
        # Remove first-person pronouns
        cleaned = self._remove_first_person_pronouns(cleaned)
        # Remove redundant duplicated sentences
        cleaned = self._remove_redundant_sentences(cleaned)
        # Collapse double spaces
        cleaned = re.sub(r'\s{2,}', ' ', cleaned)
        return cleaned.strip()

    def _remove_first_person_pronouns(self, text: str) -> str:
        """
        Remove first-person pronouns (I, we, my, our) to keep statements impersonal.
        """
        if not text:
            return text
        # Remove leading pronouns
        text = re.sub(r'^(?:I|We|My|Our)\s+', '', text, flags=re.IGNORECASE)
        # Remove standalone pronouns within the text
        text = re.sub(r'\b(?:I|We|My|Our)\b\s*', '', text, flags=re.IGNORECASE)
        return text.strip()

    def _remove_redundant_sentences(self, text: str) -> str:
        """
        Remove highly similar or alternative-form sentences from AI output.
        Keeps the first sentence and discards later ones if they are near duplicates
        or explicitly marked as alternatives.
        """
        if not text:
            return text

        sentences = re.split(r'(?<=[.!?])\s+', text.strip())
        if len(sentences) <= 1:
            return text.strip()

        def normalize(sentence: str) -> List[str]:
            sentence = sentence.strip()
            sentence = re.sub(r'^(or|alternatively|or,\s*alternatively)\s*[:,]?\s*', '', sentence, flags=re.IGNORECASE)
            words = re.findall(r'\b\w+\b', sentence.lower())
            return words

        kept_sentences = []
        kept_vectors = []
        for sentence in sentences:
            stripped = sentence.strip()
            if not stripped:
                continue
            words = normalize(stripped)
            if not words:
                continue
            is_redundant = False
            for existing_words in kept_vectors:
                overlap = set(words) & set(existing_words)
                union = set(words) | set(existing_words)
                if union and (len(overlap) / len(union)) >= 0.7:
                    is_redundant = True
                    break
            if not is_redundant:
                kept_sentences.append(stripped.rstrip(' \t'))
                kept_vectors.append(words)

        return ' '.join(kept_sentences).strip()
    
    def enhance_task(self, task: str) -> str:
        """
        Enhance a task description using available AI service (Ollama, HuggingFace, or fallback)
        Converts to past tense and makes it more professional
        
        Args:
            task: Original task description
            
        Returns:
            Enhanced task in past tense
        """
        # Try Ollama first (offline mode)
        if self.use_ollama:
            enhanced = self.enhance_task_with_ollama(task)
            if enhanced:
                return enhanced
        
        # Try HuggingFace if Ollama didn't work or isn't enabled
        enhanced = self.enhance_task_with_huggingface(task)
        if enhanced:
            return enhanced
        
        # Fallback to simple conversion
        return self._simple_past_tense_conversion(task)
    
    def _simple_past_tense_conversion(self, task: str) -> str:
        """
        Simple past tense conversion (fallback method)
        Converts common verbs to past tense and ensures professional formatting
        
        Args:
            task: Task description
            
        Returns:
            Task in past tense
        """
        task = task.strip()
        if not task:
            return task
        
        # Common verb conversions (present -> past)
        conversions = {
            r'^do\s+': 'did ',
            r'^does\s+': 'did ',
            r'^create\s+': 'created ',
            r'^develop\s+': 'developed ',
            r'^implement\s+': 'implemented ',
            r'^complete\s+': 'completed ',
            r'^finish\s+': 'finished ',
            r'^work on\s+': 'worked on ',
            r'^work\s+': 'worked ',
            r'^assist\s+': 'assisted ',
            r'^participate\s+': 'participated ',
            r'^attend\s+': 'attended ',
            r'^prepare\s+': 'prepared ',
            r'^review\s+': 'reviewed ',
            r'^update\s+': 'updated ',
            r'^fix\s+': 'fixed ',
            r'^add\s+': 'added ',
            r'^modify\s+': 'modified ',
            r'^coordinate\s+': 'coordinated ',
            r'^support\s+': 'supported ',
            r'^deliver\s+': 'delivered ',
            r'^document\s+': 'documented ',
            r'^analyze\s+': 'analyzed ',
            r'^design\s+': 'designed ',
            r'^test\s+': 'tested ',
            r'^deploy\s+': 'deployed ',
            r'^write\s+': 'wrote ',
            r'^build\s+': 'built ',
            r'^meet\s+': 'met ',
            r'^discuss\s+': 'discussed ',
            r'^present\s+': 'presented ',
            r'^organize\s+': 'organized ',
            r'^manage\s+': 'managed ',
            r'^lead\s+': 'led ',
            r'^conduct\s+': 'conducted ',
            r'^perform\s+': 'performed ',
            r'^execute\s+': 'executed ',
        }
        
        # Check if already in past tense
        past_tense_patterns = [
            r'^(completed|finished|worked|assisted|participated|attended|prepared|reviewed|updated|fixed|added|modified|coordinated|supported|delivered|documented|analyzed|designed|tested|deployed|created|developed|implemented|wrote|built|met|discussed|presented|organized|managed|led|conducted|performed|executed)\s+',
        ]
        
        task_lower = task.lower()
        is_past_tense = any(re.match(pattern, task_lower) for pattern in past_tense_patterns)
        
        if not is_past_tense:
            # Try to convert first verb
            converted = False
            for pattern, replacement in conversions.items():
                if re.match(pattern, task_lower):
                    task = re.sub(pattern, replacement, task, flags=re.IGNORECASE)
                    converted = True
                    break
            
            # If no conversion happened and doesn't start with past tense, add a verb
            if not converted and not any(task_lower.startswith(v) for v in ['completed', 'finished', 'worked', 'assisted', 'participated']):
                # Try to make it more professional by ensuring it starts with an action verb
                if not task_lower[0].isupper():
                    task = task[0].upper() + task[1:] if len(task) > 1 else task.upper()
        
        # Ensure proper capitalization
        if task:
            task = task[0].upper() + task[1:] if len(task) > 1 else task.upper()
        
        # Ensure it ends with a period if it's a complete sentence
        if task and not task.endswith(('.', '!', '?')):
            # Only add period if it looks like a complete sentence
            if len(task.split()) > 3:
                task += '.'
        
        return task
    
    def distribute_tasks_to_weeks(self, 
                                  tasks: List[str], 
                                  weeks: List[Dict],
                                  default_year: int,
                                  default_month: int) -> Dict[str, List[str]]:
        """
        Distribute tasks evenly across weeks
        
        Args:
            tasks: List of task descriptions
            weeks: List of week dictionaries
            default_year: Year of the reporting period (used when task dates omit year)
            default_month: Month of the reporting period (used when task dates omit month)
            
        Returns:
            Dictionary mapping week keys to task lists
        """
        if not weeks:
            return {}
        
        # Enhance all tasks
        enhanced_tasks = [self.enhance_task(task) for task in tasks]
        week_keys = [self.format_week_key(week) for week in weeks]
        week_tasks = {key: [] for key in week_keys}

        unassigned_tasks = []
        for original_task, enhanced_task in zip(tasks, enhanced_tasks):
            date_ranges = self._parse_task_date_ranges(original_task, default_year, default_month)
            week_index = -1
            for date_range in date_ranges:
                week_index = self._find_week_index_for_range(date_range, weeks)
                if week_index != -1:
                    break
            if week_index != -1:
                week_key = week_keys[week_index]
                week_tasks[week_key].append(enhanced_task)
            else:
                unassigned_tasks.append(enhanced_task)

        # Distribute remaining tasks evenly
        if unassigned_tasks:
            tasks_per_week = len(unassigned_tasks) // len(weeks)
            remainder = len(unassigned_tasks) % len(weeks)
            task_index = 0
            for i, week_key in enumerate(week_keys):
                num_tasks = tasks_per_week + (1 if i < remainder else 0)
                if num_tasks > 0:
                    week_tasks[week_key].extend(unassigned_tasks[task_index:task_index + num_tasks])
                    task_index += num_tasks

        return week_tasks
    
    def format_date_range(self, start: datetime, end: datetime) -> str:
        """
        Format date range as "Month Day-Day, Year"
        
        Args:
            start: Start date
            end: End date
            
        Returns:
            Formatted date string
        """
        if start.date() == end.date():
            return start.strftime('%B %d, %Y')
        if start.month == end.month:
            return f"{start.strftime('%B %d')}-{end.strftime('%d, %Y')}"
        else:
            return f"{start.strftime('%B %d')}-{end.strftime('%B %d, %Y')}"
    
    def format_week_key(self, week: Dict) -> str:
        """
        Format week dictionary as "Month Day-Day, Year"
        
        Args:
            week: Week dictionary with 'start' and 'end' keys
            
        Returns:
            Formatted week date string
        """
        return self.format_date_range(week['start'], week['end'])
    
    def generate_report(self, 
                       name: str,
                       position: str,
                       office: str,
                       year: int,
                       month: int,
                       is_first_half: bool,
                       tasks: List[str],
                       reviewed_by: str = "",
                       verified_by: str = "",
                       approved_by: str = "",
                       accepted_by: str = "",
                       output_filename: str = None) -> str:
        """
        Generate the accomplishment report document
        
        Args:
            name: Employee name
            position: Employee position
            office: Office name
            year: Report year
            month: Report month (1-12)
            is_first_half: True for 1-15, False for 16-end
            tasks: List of task descriptions
            reviewed_by: Name and details of reviewer (optional)
            verified_by: Name and details of verifier (optional)
            approved_by: Name and details of approver (optional)
            accepted_by: Name and details of acceptor (optional)
            output_filename: Output filename (auto-generated if None)
            
        Returns:
            Path to generated document
        """
        # Calculate date ranges
        start_first, end_first, start_second, end_second, last_day = self.calculate_date_ranges(year, month)
        
        if is_first_half:
            period_start = start_first
            period_end = end_first
        else:
            period_start = start_second
            period_end = end_second
        
        # Get working days and split into weeks
        working_days = self.get_working_days(period_start, period_end)
        weeks = self.split_into_weeks(working_days)
        
        # Distribute tasks to weeks (respect explicit dates when present)
        week_tasks = self.distribute_tasks_to_weeks(
            tasks,
            weeks,
            period_start.year,
            period_start.month
        )
        
        # Create document
        doc = Document()
        
        # Set margins to match reference (1 inch on all sides)
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
        
        # Add header
        header_para = doc.add_paragraph('INDIVIDUAL ACCOMPLISHMENT REPORT')
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.runs[0]
        header_run.font.size = Pt(14)
        header_run.bold = True
        
        doc.add_paragraph()  # Empty line
        
        # Add employee info
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('NAME: ')
        name_run.bold = True
        name_para.add_run(f'{name}')

        position_para = doc.add_paragraph()
        position_run = position_para.add_run('POSITION: ')
        position_run.bold = True
        position_para.add_run(f'{position}')

        office_para = doc.add_paragraph()
        office_run = office_para.add_run('OFFICE: ')
        office_run.bold = True
        office_para.add_run(f'{office}')

        # Add date
        date_str = self.format_date_range(period_start, period_end)
        date_para = doc.add_paragraph()
        date_run = date_para.add_run('DATE: ')
        date_run.bold = True
        date_para.add_run(f'{date_str}')
        
        doc.add_paragraph()  # Empty line
        
        # Create table with black borders
        table = doc.add_table(rows=1, cols=2)
        table.style = None  # No style, we'll add black borders manually
        
        # Set table-level borders to black
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tbl.insert(0, tblPr)
        
        # Set black borders for table
        def set_table_borders(tblPr):
            """Set table borders to black"""
            
            borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is None:
                tblBorders = OxmlElement('w:tblBorders')
                tblPr.append(tblBorders)
            else:
                # Clear existing borders
                for border in list(tblBorders):
                    tblBorders.remove(border)
            
            for border_name in borders:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')  # Thicker border (12 = 1.5pt)
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Black
                tblBorders.append(border)
        
        set_table_borders(tblPr)
        
        # Set black borders for all cells
        def set_black_borders(cell):
            """Set all borders of a cell to black - ensures visible borders"""
            
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Create border elements for all sides
            borders = ['top', 'left', 'bottom', 'right']
            for border_name in borders:
                tag = qn(f'w:{border_name}')
                # Remove existing border if any
                existing = tcPr.find(tag)
                if existing is not None:
                    tcPr.remove(existing)
                
                # Create new border element with thicker, visible black border
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')  # Thicker border (12 = 1.5pt)
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Black - ensure it's explicitly set
                tcPr.append(border)
        
        # Set column widths - PERIOD/WEEK column narrower, ACCOMPLISHMENT/OUTPUT wider
        table.columns[0].width = Inches(1.5)  # Narrow column for dates
        table.columns[1].width = Inches(5.5)  # Wide column for accomplishments
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'PERIOD/ WEEK'
        header_cells[1].text = 'ACCOMPLISHMENT / OUTPUT'
        
        # Make header bold and set black borders
        for cell in header_cells:
            set_black_borders(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add week rows
        for week in weeks:
            week_key = self.format_week_key(week)
            tasks_for_week = week_tasks.get(week_key, [])
            
            if tasks_for_week:
                row = table.add_row()
                row.cells[0].text = week_key
                
                # Format tasks with Word bullet style for consistent rendering
                tasks_cell = row.cells[1]
                # Clear existing empty paragraph
                tasks_cell.text = ''
                if tasks_cell.paragraphs:
                    p = tasks_cell.paragraphs[0]
                    p._element.getparent().remove(p._element)
                
                for task in tasks_for_week:
                    para = tasks_cell.add_paragraph(task)
                    para.style = 'List Bullet'
                
                # Set black borders for all cells in this row
                for cell in row.cells:
                    set_black_borders(cell)
        
        # Add signature section
        doc.add_paragraph()  # Empty line
        doc.add_paragraph()  # Empty line
        
        # Create signature table (no borders)
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.style = None  # No style, no borders
        
        # Signature table - no borders needed
        sig_tbl = sig_table._tbl
        sig_tblPr = sig_tbl.tblPr
        if sig_tblPr is None:
            sig_tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            sig_tbl.insert(0, sig_tblPr)
        # Don't set borders for signature table - it should have no borders
        
        sig_cells = sig_table.rows[0].cells
        
        # Helper function to format signature text with proper font sizes
        def format_signature_cell(cell, prepared_text, verified_text, accepted_text):
            """Format signature cell with different font sizes"""
            cell.text = ''  # Clear cell
            para = cell.paragraphs[0]
            
            # Prepared by: (9pt, bold)
            run = para.add_run('Prepared by:')
            run.font.size = Pt(9)
            run.bold = True
            para.add_run('\n\n\n')
            
            # Name (10pt)
            prepared_lines = prepared_text.split('\n')
            if prepared_lines:
                run = para.add_run(prepared_lines[0])
                run.font.size = Pt(10)
                run.bold = True
                # Details (8pt) - position and office
                if len(prepared_lines) > 1:
                    for detail_line in prepared_lines[1:]:
                        para.add_run('\n')
                        run = para.add_run(detail_line)
                        run.font.size = Pt(8)
            
            para.add_run('\n\n\n\n')
            
            # Verified by: (9pt, bold)
            run = para.add_run('Verified by:')
            run.font.size = Pt(9)
            run.bold = True
            
            if verified_text:
                para.add_run('\n\n\n')
                # Split verified text into name and details
                verified_lines = verified_text.split('\n')
                if verified_lines:
                    run = para.add_run(verified_lines[0])
                    run.font.size = Pt(10)
                    run.bold = True
                    if len(verified_lines) > 1:
                        for detail_line in verified_lines[1:]:
                            para.add_run('\n')
                            run = para.add_run(detail_line)
                            run.font.size = Pt(8)
            else:
                para.add_run('\n\n\n\n\n\n\n\n\n')
            
            para.add_run('\n\n\n\n')
            
            # Accepted by: (9pt, bold)
            run = para.add_run('Accepted by: ')
            run.font.size = Pt(9)
            run.bold = True
            
            if accepted_text:
                para.add_run('\n\n\n')
                # Split accepted text into name and details
                accepted_lines = accepted_text.split('\n')
                if accepted_lines:
                    run = para.add_run(accepted_lines[0])
                    run.font.size = Pt(10)
                    run.bold = True
                    if len(accepted_lines) > 1:
                        for detail_line in accepted_lines[1:]:
                            para.add_run('\n')
                            run = para.add_run(detail_line)
                            run.font.size = Pt(8)
            else:
                para.add_run('\n\n\n\n\n\n\n\n\n')
        
        def format_signature_cell_right(cell, reviewed_text, approved_text):
            """Format right signature cell with proper font sizes"""
            cell.text = ''  # Clear cell
            para = cell.paragraphs[0]
            
            # Reviewed by: (9pt, bold)
            run = para.add_run('Reviewed by:')
            run.font.size = Pt(9)
            run.bold = True
            
            if reviewed_text:
                para.add_run('\n\n\n')
                # Split reviewed text into name and details
                reviewed_lines = reviewed_text.split('\n')
                if reviewed_lines:
                    run = para.add_run(reviewed_lines[0])
                    run.font.size = Pt(10)
                    run.bold = True
                    if len(reviewed_lines) > 1:
                        for detail_line in reviewed_lines[1:]:
                            para.add_run('\n')
                            run = para.add_run(detail_line)
                            run.font.size = Pt(8)
            else:
                para.add_run('\n\n\n\n\n\n\n\n\n')
            
            para.add_run('\n\n\n\n')
            
            # Approved by: (9pt, bold)
            run = para.add_run('Approved by:')
            run.font.size = Pt(9)
            run.bold = True
            
            if approved_text:
                para.add_run('\n\n\n')
                # Split approved text into name and details
                approved_lines = approved_text.split('\n')
                if approved_lines:
                    run = para.add_run(approved_lines[0])
                    run.font.size = Pt(10)
                    run.bold = True
                    if len(approved_lines) > 1:
                        for detail_line in approved_lines[1:]:
                            para.add_run('\n')
                            run = para.add_run(detail_line)
                            run.font.size = Pt(8)
            else:
                para.add_run('\n\n\n\n\n\n\n\n\n')
        
        # Handle line breaks in signature fields (replace \n with actual newlines)
        verified_by_processed = verified_by.replace('\\n', '\n') if verified_by else ''
        accepted_by_processed = accepted_by.replace('\\n', '\n') if accepted_by else ''
        reviewed_by_processed = reviewed_by.replace('\\n', '\n') if reviewed_by else ''
        approved_by_processed = approved_by.replace('\\n', '\n') if approved_by else ''
        
        # Format left cell
        prepared_text = f'{name}\n{position}\n{office}'
        format_signature_cell(sig_cells[0], prepared_text, verified_by_processed, accepted_by_processed)
        
        # Format right cell
        format_signature_cell_right(sig_cells[1], reviewed_by_processed, approved_by_processed)
        
        # Remove borders from signature table - make them transparent/no border
        for cell in sig_cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Remove all borders
            for border_name in ['top', 'left', 'bottom', 'right']:
                tag = qn(f'w:{border_name}')
                existing = tcPr.find(tag)
                if existing is not None:
                    tcPr.remove(existing)
        
        # Also remove table-level borders for signature table
        sig_tblBorders = sig_tblPr.find(qn('w:tblBorders'))
        if sig_tblBorders is not None:
            sig_tblPr.remove(sig_tblBorders)
        
        # Generate output filename
        if not output_filename:
            month_name = calendar.month_name[month]
            output_filename = f"ACCOMPLISHMENT_REPORT_{month_name}_{period_start.day}-{period_end.day}_{year}.docx"
        
        # Save document
        doc.save(output_filename)
        
        return output_filename


def read_config_from_file(filename: str) -> Dict:
    """
    Read configuration and tasks from a text file
    
    Args:
        filename: Path to the input file
        
    Returns:
        Dictionary with keys: name, position, office, year, month, period, api_key, 
        reviewed_by, verified_by, approved_by, accepted_by, tasks
    """
    config = {
        'name': None,
        'position': None,
        'office': None,
        'year': None,
        'month': None,
        'period': None,
        'api_key': None,
        'reviewed_by': None,
        'verified_by': None,
        'approved_by': None,
        'accepted_by': None,
        'use_ollama': False,
        'ollama_model': 'llama3.1',
        'ollama_base_url': 'http://localhost:11434',
        'tasks': []
    }
    
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        tasks_section = False
        for line in lines:
            original_line = line
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Skip comments
            if line.startswith('#'):
                # Check if we're entering tasks section (look for section marker)
                if '===== TASKS =====' in line.upper() or line.upper().strip() == '# TASKS':
                    tasks_section = True
                continue
            
            # Parse key-value pairs (only if not in tasks section)
            if ':' in line and not tasks_section:
                key, value = line.split(':', 1)
                key = key.strip().upper()
                value = value.strip()
                
                if key == 'NAME':
                    config['name'] = value
                elif key == 'POSITION':
                    config['position'] = value
                elif key == 'OFFICE':
                    config['office'] = value
                elif key == 'YEAR':
                    try:
                        config['year'] = int(value)
                    except ValueError:
                        pass
                elif key == 'MONTH':
                    try:
                        config['month'] = int(value)
                    except ValueError:
                        pass
                elif key == 'PERIOD':
                    try:
                        config['period'] = int(value)
                    except ValueError:
                        pass
                elif key == 'API_KEY':
                    config['api_key'] = value if value else None
                elif key == 'REVIEWED_BY':
                    config['reviewed_by'] = value
                elif key == 'VERIFIED_BY':
                    config['verified_by'] = value
                elif key == 'APPROVED_BY':
                    config['approved_by'] = value
                elif key == 'ACCEPTED_BY':
                    config['accepted_by'] = value
                elif key == 'USE_OLLAMA':
                    config['use_ollama'] = value.lower() in ('true', 'yes', '1', 'y')
                elif key == 'OLLAMA_MODEL':
                    config['ollama_model'] = value
                elif key == 'OLLAMA_BASE_URL':
                    config['ollama_base_url'] = value
                # Don't treat key-value pairs as tasks
                continue
            
            # If we're in tasks section or past the config section, add as task
            if tasks_section:
                # We're in the tasks section, add as task
                config['tasks'].append(line)
            # Only treat as task if it's not a key-value pair and we've seen the TASKS section marker
            # (This handles old format where tasks come after config)
        
    except FileNotFoundError:
        print(f"Error: File '{filename}' not found.")
    except Exception as e:
        print(f"Error reading file: {e}")
    
    return config


def main():
    """Main function for command-line usage"""
    print("=" * 80)
    print("ACCOMPLISHMENT REPORT GENERATOR")
    print("=" * 80)
    print()
    
    # Ask if user wants to read from file
    use_file = input("Read from input file? (y/n, default: y): ").strip().lower()
    if not use_file:
        use_file = 'y'
    
    config = {}
    filename = "tasks_input.txt"
    
    if use_file == 'y':
        filename_input = input(f"Enter filename (default: {filename}): ").strip()
        if filename_input:
            filename = filename_input
        
        config = read_config_from_file(filename)
        
        if config.get('tasks'):
            print(f"\n✓ Loaded configuration from {filename}")
            if config.get('name'):
                print(f"  Name: {config['name']}")
            if config.get('position'):
                print(f"  Position: {config['position']}")
            if config.get('office'):
                print(f"  Office: {config['office']}")
            if config.get('year') and config.get('month'):
                period_str = "1-15" if config.get('period') == 1 else "16-end"
                print(f"  Period: {config['year']}, Month {config['month']}, {period_str}")
            print(f"  Tasks: {len(config['tasks'])} tasks loaded")
            print()
            
            # Show tasks
            show_tasks = input("Show tasks? (y/n, default: n): ").strip().lower()
            if show_tasks == 'y':
                for i, task in enumerate(config['tasks'], 1):
                    print(f"  {i}. {task}")
                print()
    
    # Get user information - use from file if available, otherwise prompt
    name = config.get('name')
    if not name:
        name = input("Enter your name: ").strip()
    
    position = config.get('position')
    if not position:
        position = input("Enter your position: ").strip()
    
    office = config.get('office')
    if not office:
        office = input("Enter your office: ").strip()
    
    # Get date information - use from file if available, otherwise prompt
    year = config.get('year')
    if not year:
        year = int(input("Enter year (e.g., 2025): "))
    
    month = config.get('month')
    if not month:
        month = int(input("Enter month (1-12): "))
    
    period = config.get('period')
    if not period:
        half = input("Enter period (1 for 1-15, 2 for 16-end): ").strip()
        is_first_half = half == "1"
    else:
        is_first_half = period == 1
    
    # Get tasks
    tasks = config.get('tasks', [])
    if not tasks:
        print("\nEnter your tasks (one per line, press Enter twice when done):")
        while True:
            task = input()
            if not task.strip():
                if tasks:  # If we have tasks and user pressed enter again, we're done
                    break
                continue
            tasks.append(task.strip())
    
    if not tasks:
        print("Error: No tasks provided. Exiting.")
        return
    
    # Get HuggingFace API key (optional) - use from file if available
    api_key = config.get('api_key')
    if not api_key:
        api_key = input("\nEnter HuggingFace API key (optional, press Enter to skip): ").strip()
        if not api_key:
            api_key = None
    
    # Get Ollama configuration - use from file if available
    use_ollama = config.get('use_ollama', False)
    ollama_model = config.get('ollama_model', 'llama3.1')
    ollama_base_url = config.get('ollama_base_url', 'http://localhost:11434')
    
    if not api_key and not use_ollama:
        print("Using simple past tense conversion (no API enhancement)")
    elif use_ollama:
        print(f"Using Ollama for task enhancement (model: {ollama_model})")
    elif api_key:
        print("Using HuggingFace API for task enhancement")
    
    # Get signature information - use from file if available
    reviewed_by = config.get('reviewed_by', '')
    verified_by = config.get('verified_by', '')
    approved_by = config.get('approved_by', '')
    accepted_by = config.get('accepted_by', '')
    
    # Generate report
    generator = AccomplishmentGenerator(
        huggingface_api_key=api_key,
        use_ollama=use_ollama,
        ollama_model=ollama_model,
        ollama_base_url=ollama_base_url
    )
    
    print("\nGenerating report...")
    output_file = generator.generate_report(
        name=name,
        position=position,
        office=office,
        year=year,
        month=month,
        is_first_half=is_first_half,
        tasks=tasks,
        reviewed_by=reviewed_by,
        verified_by=verified_by,
        approved_by=approved_by,
        accepted_by=accepted_by
    )
    
    print(f"\nReport generated successfully: {output_file}")


if __name__ == "__main__":
    main()

